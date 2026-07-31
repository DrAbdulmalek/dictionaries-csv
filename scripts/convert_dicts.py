#!/usr/bin/env python3
"""
Convert every dictionary file in the extracted folder to its own CSV.

Schema (UTF-8 CSV with header):
    term, definition, source_file

Rules:
- One CSV per source dictionary (no merging across files).
- File name = source dict name + ".csv".
- For unsupported formats (BGL/LD2), write an empty CSV with a NOTE row.

Supported formats:
- StarDict   : .ifo + .idx + .dict (and .dict.dz / .syn.dz)
- MDict      : .mdx (+ optional .mdd for resources)
- HTML dicts : .html (gzipped or plain) — extract <b> term + <div> definition
- PDF        : text-based PDFs (pdfplumber)
- DOCX       : python-docx, one row per paragraph
"""

from __future__ import annotations

import csv
import gzip
import io
import os
import re
import sys
import struct
import traceback
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Iterator, Optional

# ─── Bootstrap: stub lzo so readmdict imports without system lzo ───
_lzo_stub = type(sys)("lzo")
def _safe_lzo_decompress(data, unused=None):
    """Detect real LZO compression and refuse to silently corrupt data."""
    if data[:4] == b"\x89LZO":
        raise RuntimeError("Real LZO compression detected. Install: pip install python-lzo")
    return data

_lzo_stub.decompress = _safe_lzo_decompress
_lzo_stub.compress = lambda data, unused=None: data
sys.modules.setdefault("lzo", _lzo_stub)

# ─── Now safe to import the heavy deps ───
import bs4  # noqa: E402
from bs4 import BeautifulSoup  # noqa: E402

import argparse as _argparse
_parser = _argparse.ArgumentParser(add_help=False)
_parser.add_argument("--root", default=os.environ.get("DICT_WORK_DIR", "."))
_args, _ = _parser.parse_known_args()
ROOT = Path(_args.root).resolve()
EXTRACTED = ROOT / "extracted" / "New Folder"
UNPACKED = ROOT / "unpacked"
CSV_OUT = ROOT / "csv_output"
LOGS = ROOT / "logs"
CSV_OUT.mkdir(parents=True, exist_ok=True)
LOGS.mkdir(parents=True, exist_ok=True)

# ─── Helpers ────────────────────────────────────────────────────────


def _clean(text: str) -> str:
    """Collapse whitespace, strip, but preserve Arabic characters."""
    if text is None:
        return ""
    # Replace NBSP and other weird whitespace
    text = text.replace("\xa0", " ").replace("\u200b", "").replace("\u200f", "")
    # Collapse runs of whitespace (including newlines) into a single space
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _strip_html(html: str) -> str:
    """Convert HTML to plain text, preserving a single space between blocks."""
    if not html:
        return ""
    soup = BeautifulSoup(html, "lxml")
    # Remove scripts/styles
    for tag in soup(["script", "style", "link"]):
        tag.decompose()
    return _clean(soup.get_text(separator=" "))


def _safe_csv_name(name: str) -> str:
    """Make a filename safe for the CSV output."""
    safe = re.sub(r"[^\w\u0600-\u06FF\u0750-\u077F\-. ]+", "_", name)
    safe = re.sub(r"\.{2,}", ".", safe)
    return safe.strip().rstrip("._")


def _write_csv(rows: list[tuple[str, str, str]], out_path: Path, source_name: str) -> int:
    """Write rows to a UTF-8 CSV (with BOM for Excel compatibility)."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f, quoting=csv.QUOTE_MINIMAL)
        w.writerow(["term", "definition", "source_file"])
        for term, definition, _ in rows:
            w.writerow([term, definition, source_name])
    return len(rows)


def _note_csv(out_path: Path, source_name: str, note: str) -> None:
    """Write a single-row CSV noting that conversion was not possible."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with open(out_path, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow(["term", "definition", "source_file"])
        w.writerow(["[CONVERSION NOTE]", note, source_name])


# ─── StarDict converter ────────────────────────────────────────────


def _read_dict_dz(path: Path) -> bytes:
    """Read a .dict file, transparently decompressing .dict.dz if present."""
    if path.with_suffix(path.suffix + ".dz").exists():   # x.dict -> x.dict.dz
        with gzip.open(path.with_suffix(path.suffix + ".dz"), "rb") as f:
            return f.read()
    if path.suffix == ".dz":
        with gzip.open(path, "rb") as f:
            return f.read()
    return path.read_bytes()


def parse_stardict(ifo_path: Path) -> list[tuple[str, str, str]]:
    """Parse StarDict (.ifo + .idx + .dict) and return list of (term, definition, source)."""
    base = ifo_path.with_suffix("")   # remove .ifo
    idx_path = base.with_suffix(".idx")
    dict_path = base.with_suffix(".dict")

    # Handle .dict.dz: the base dict file might be x.dict (and x.dict.dz is the compressed version)
    if not dict_path.exists():
        if dict_path.with_suffix(".dict.dz").exists():
            dict_data = _read_dict_dz(dict_path.with_suffix(".dict"))
        else:
            raise FileNotFoundError(f"Missing .dict for {ifo_path}")
    else:
        dict_data = _read_dict_dz(dict_path)

    # Parse .ifo for metadata (sametypesequence etc.)
    ifo_text = ifo_path.read_text(encoding="utf-8", errors="replace")
    bookname = "Unknown"
    sametypesequence = None
    for line in ifo_text.splitlines():
        if line.startswith("bookname="):
            bookname = line.split("=", 1)[1].strip()
        elif line.startswith("sametypesequence="):
            sametypesequence = line.split("=", 1)[1].strip()

    # Parse .idx: each entry is: word\0 + 4-byte big-endian offset + 4-byte big-endian size
    idx_data = idx_path.read_bytes()
    entries: list[tuple[str, str, str]] = []
    i = 0
    n = len(idx_data)
    while i < n:
        # Find null terminator
        end = idx_data.find(b"\x00", i)
        if end == -1:
            break
        word = idx_data[i:end].decode("utf-8", errors="replace")
        # 8 bytes after null: offset (4) + size (4), big-endian
        if end + 9 > n:
            break
        offset, size = struct.unpack(">II", idx_data[end + 1: end + 9])
        definition_bytes = dict_data[offset: offset + size]
        # StarDict sametypesequence 'h' means HTML; 'm' means plain text; 'x' means xdxf
        if sametypesequence and sametypesequence[0] in ("h", "x"):
            definition = _strip_html(definition_bytes.decode("utf-8", errors="replace"))
        else:
            definition = _clean(definition_bytes.decode("utf-8", errors="replace"))
        entries.append((word, definition, ifo_path.stem))
        i = end + 9
    return entries


# ─── MDict (.mdx) converter ────────────────────────────────────────


def parse_mdx(mdx_path: Path) -> list[tuple[str, str, str]]:
    """Parse MDict MDX file and return list of (term, definition, source)."""
    from readmdict import MDX
    mdx = MDX(str(mdx_path))
    entries: list[tuple[str, str, str]] = []
    for key, value in mdx.items():
        term = key.decode("utf-8", errors="replace").strip()
        if not term:
            continue
        # MDict entries are HTML
        definition = _strip_html(value.decode("utf-8", errors="replace"))
        entries.append((term, definition, mdx_path.stem))
    return entries


# ─── dicthtml (gzipped HTML) converter ─────────────────────────────


def parse_dicthtml_folder(folder: Path) -> list[tuple[str, str, str]]:
    """Parse a folder of (optionally gzipped) HTML dictionary files."""
    entries: list[tuple[str, str, str]] = []
    for html_file in sorted(folder.glob("*.html")):
        raw = html_file.read_bytes()
        if raw[:2] == b"\x1f\x8b":
            try:
                raw = gzip.decompress(raw)
            except Exception:
                continue
        try:
            soup = BeautifulSoup(raw, "lxml")
        except Exception:
            continue
        # dicthtml-en-ar uses <w> elements, each with <a name="TERM"/> + <div><b>TERM</b><br/>DEFINITION</div>
        for w in soup.find_all("w"):
            a = w.find("a")
            term = ""
            if a and a.get("name"):
                term = a["name"]
            else:
                b = w.find("b")
                if b:
                    term = b.get_text(strip=True)
            if not term:
                continue
            # Definition = the text inside <div> after <b>
            div = w.find("div")
            if div:
                # Remove <b> from div to get definition only
                b = div.find("b")
                if b:
                    b.extract()
                definition = _clean(div.get_text(separator=" "))
            else:
                definition = _clean(w.get_text(separator=" "))
            entries.append((term, definition, html_file.stem))
    return entries


# ─── PDF converter ─────────────────────────────────────────────────


def parse_pdf(pdf_path: Path) -> list[tuple[str, str, str]]:
    """Parse a PDF dictionary. Heuristic: detect two-column layouts (EN | AR)."""
    import pdfplumber
    entries: list[tuple[str, str, str]] = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page in pdf.pages:
            # Try table extraction first (most dictionaries are table-like)
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    for row in table:
                        cells = [c for c in (row or []) if c and c.strip()]
                        if not cells:
                            continue
                        if len(cells) >= 2:
                            term = _clean(cells[0])
                            definition = _clean(" ".join(cells[1:]))
                        else:
                            single = _clean(cells[0])
                            term, definition = single, ""
                        if term or definition:
                            entries.append((term, definition, pdf_path.stem))
            else:
                # No tables — fall back to line-by-line text
                text = page.extract_text() or ""
                for line in text.splitlines():
                    line = _clean(line)
                    if not line:
                        continue
                    # Heuristic: split on 2+ spaces OR tab (typical bilingual layout)
                    parts = re.split(r"\s{2,}|\t", line, maxsplit=1)
                    if len(parts) == 2:
                        entries.append((parts[0], parts[1], pdf_path.stem))
                    else:
                        entries.append((line, "", pdf_path.stem))
    return entries


# ─── DOCX converter ────────────────────────────────────────────────


def parse_docx(docx_path: Path) -> list[tuple[str, str, str]]:
    """Parse a DOCX file: each non-empty paragraph becomes one row."""
    from docx import Document
    doc = Document(str(docx_path))
    entries: list[tuple[str, str, str]] = []
    # Also try tables first
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if not cells:
                continue
            if len(cells) >= 2:
                entries.append((_clean(cells[0]), _clean(" ".join(cells[1:])), docx_path.stem))
            else:
                entries.append((_clean(cells[0]), "", docx_path.stem))
    # Then paragraphs (in case there are no tables)
    if not entries:
        for para in doc.paragraphs:
            text = _clean(para.text)
            if not text:
                continue
            # Heuristic: split on tab or 2+ spaces
            parts = re.split(r"\t|\s{2,}", text, maxsplit=1)
            if len(parts) == 2:
                entries.append((parts[0], parts[1], docx_path.stem))
            else:
                entries.append((text, "", docx_path.stem))
    return entries


# ─── Dispatcher ────────────────────────────────────────────────────

CONVERSIONS: list[tuple[str, str, Path, callable]] = []


def register(name: str, src: Path, parser: callable, csv_name: Optional[str] = None):
    csv_name = csv_name or _safe_csv_name(name) + ".csv"
    CONVERSIONS.append((name, csv_name, src, parser))


# 1) Standalone MDict files
register("Dictionary of Medical Terms",
         EXTRACTED / "Dictionary of Medical Terms.mdx",
         parse_mdx)
register("Longman English Arabic",
         EXTRACTED / "Longman English Arabic.mdx",
         parse_mdx)
# Longman Modern En-En-Ar (from .zip and .rar are duplicates — process standalone .mdx equivalent)
register("Longman English Arabic (from RAR)",
         UNPACKED / "Longman_Modern_EnEnAr_rar" / "Longman Modern En-En-Ar" / "Longman English Arabic.mdx",
         parse_mdx,
         csv_name="Longman_English_Arabic_from_RAR.csv")
register("Longman English Arabic (from ZIP)",
         UNPACKED / "Longman_Modern_EnEnAr_zip" / "Longman Modern En-En-Ar" / "Longman English Arabic.mdx",
         parse_mdx,
         csv_name="Longman_English_Arabic_from_ZIP.csv")

# 2) Standalone StarDict
register("Longman Modern En-En-Ar",
         EXTRACTED / "Longman Modern En-En-Ar.ifo",
         parse_stardict)

# 3) Oxford dictionaries (StarDict, .dict.dz compressed)
register("Oxford Arabic Dictionary",
         UNPACKED / "Oxford_Arabic_Dictionary_1" / "Oxford_Arabic_Dictionary" / "Oxford_Arabic_Dictionary.ifo",
         parse_stardict)
register("Oxford Arabic Dictionary (En-Ar)",
         UNPACKED / "Oxford_Arabic_Dictionary_EnAr" / "Oxford Arabic Dictionary" / "Bodydata.ifo",
         parse_stardict,
         csv_name="Oxford_Arabic_Dictionary_EnAr.csv")

# 4) dicthtml-en-ar (folder of gzipped HTML files)
register("dicthtml-en-ar",
         UNPACKED / "dicthtml_en_ar",
         parse_dicthtml_folder)

# 5) PDFs
register("downloadfile-1",
         EXTRACTED / "downloadfile-1.pdf",
         parse_pdf)
register("msf-glossary",
         EXTRACTED / "msf-glossary.pdf",
         parse_pdf)
register("المصطلحات_الأدبية_الحديثة",
         EXTRACTED / "المصطلحات_الأدبية_الحديثة.pdf",
         parse_pdf)
register("معجم المصطلحات الاعلامية",
         EXTRACTED / "معجم المصطلحات الاعلامية.pdf",
         parse_pdf)
register("معجم مصطلحات الإعلام",
         EXTRACTED / "معجم مصطلحات الإعلام.pdf",
         parse_pdf)

# 6) DOCX
register("أخطاء تنقيح النصوص",
         EXTRACTED / "أخطاء_تنقيح_النصوص.docx",
         parse_docx)


# ─── Main ──────────────────────────────────────────────────────────


def main():
    summary: list[tuple[str, str, int, str]] = []   # (name, csv_name, rows, status)
    for name, csv_name, src, parser in CONVERSIONS:
        out = CSV_OUT / csv_name
        print(f"\n[convert] {name}")
        print(f"  source: {src}")
        print(f"  output: {out}")
        if not src.exists():
            print(f"  ! source not found — skipping")
            _note_csv(out, name, f"Source file not found: {src}")
            summary.append((name, csv_name, 0, "MISSING"))
            continue
        try:
            rows = parser(src)
            if not rows:
                _note_csv(out, name, "Parser returned 0 entries (file may be empty or in an unsupported sub-format).")
                summary.append((name, csv_name, 0, "EMPTY"))
                print(f"  ! 0 entries — wrote note CSV")
                continue
            count = _write_csv(rows, out, name)
            summary.append((name, csv_name, count, "OK"))
            print(f"  ✓ wrote {count} entries")
        except Exception as e:
            traceback.print_exc()
            _note_csv(out, name, f"Conversion failed: {type(e).__name__}: {e}")
            summary.append((name, csv_name, 0, f"ERROR: {type(e).__name__}"))
            print(f"  ✗ ERROR: {e}")

    # ─── BGL / LD2 (Babylon) — write note CSVs ───
    babylon_dir = UNPACKED / "ArabicDictionariesOfBabylon" / "DictionariesOfBabylon"
    if babylon_dir.exists():
        for bgl in sorted(babylon_dir.iterdir()):
            if bgl.suffix.lower() in (".bgl", ".ld2"):
                csv_name = _safe_csv_name(bgl.stem) + ".csv"
                out = CSV_OUT / csv_name
                _note_csv(
                    out,
                    bgl.name,
                    "Babylon .BGL/.LD2 is a proprietary binary format. "
                    "Use 'dictconv' (freedict) or the Babylon desktop app to export to text first, "
                    "then re-run this script on the exported text.",
                )
                summary.append((bgl.name, csv_name, 0, "UNSUPPORTED_FORMAT"))

    # ─── Summary report ───
    report_path = LOGS / "conversion_report.txt"
    with open(report_path, "w", encoding="utf-8") as f:
        f.write("=" * 90 + "\n")
        f.write("DICTIONARY → CSV CONVERSION REPORT\n")
        f.write("=" * 90 + "\n\n")
        f.write(f"{'NAME':<50} {'ROWS':>8}  {'STATUS':<22}  CSV\n")
        f.write("-" * 90 + "\n")
        for name, csv_name, rows, status in summary:
            f.write(f"{name[:48]:<50} {rows:>8}  {status:<22}  {csv_name}\n")
        f.write("\n")
        ok = sum(1 for *_, s in summary if s == "OK")
        empty = sum(1 for *_, s in summary if s == "EMPTY")
        err = sum(1 for *_, s in summary if s.startswith("ERROR"))
        missing = sum(1 for *_, s in summary if s == "MISSING")
        unsupported = sum(1 for *_, s in summary if s == "UNSUPPORTED_FORMAT")
        total_rows = sum(r for _, _, r, s in summary if s == "OK")
        f.write(f"\nTotal dictionaries: {len(summary)}\n")
        f.write(f"  OK         : {ok}\n")
        f.write(f"  EMPTY      : {empty}\n")
        f.write(f"  ERROR      : {err}\n")
        f.write(f"  MISSING    : {missing}\n")
        f.write(f"  UNSUPPORTED: {unsupported}  (Babylon BGL/LD2 — needs proprietary tools)\n")
        f.write(f"\nTotal entries extracted (OK only): {total_rows:,}\n")

    print("\n" + "=" * 70)
    print("SUMMARY")
    print("=" * 70)
    for name, csv_name, rows, status in summary:
        print(f"  [{status:<18}] {rows:>7} rows  {name[:50]:<52} -> {csv_name}")
    print(f"\nReport saved to: {report_path}")
    print(f"CSVs saved to:   {CSV_OUT}")


if __name__ == "__main__":
    main()
